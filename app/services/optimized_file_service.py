"""
Optimized File Processing Service
- Uses local libraries first (FREE) for text extraction
- Falls back to OCR for scanned PDFs (pytesseract or AI vision)
- Content hashing for deduplication
- Significant cost savings vs AI-only approach
"""

import os
import hashlib
import tempfile
from typing import Optional, Tuple, List
from app.core.config import settings

# Local libraries for FREE text extraction
try:
    import pypdf  # type: ignore
except ImportError:
    # Fall back to PyPDF2 if pypdf not installed
    import PyPDF2 as pypdf  # type: ignore
from docx import Document

# Optional OCR dependencies
try:
    from pdf2image import convert_from_path  # type: ignore
    PDF2IMAGE_AVAILABLE = True
except ImportError:
    PDF2IMAGE_AVAILABLE = False
    convert_from_path = None

try:
    import pytesseract  # type: ignore
    TESSERACT_AVAILABLE = True
except ImportError:
    TESSERACT_AVAILABLE = False
    pytesseract = None

class OptimizedFileService:
    """
    Cost-optimized file processing:
    1. Uses PyPDF2/python-docx for text extraction (FREE)
    2. Uses pytesseract OCR for scanned PDFs (FREE, requires Tesseract)
    3. Falls back to AI Vision for OCR when pytesseract unavailable
    4. Caches extracted text by file hash
    """
    
    def __init__(self):
        self._text_cache: dict = {}  # file_hash -> extracted_text
    
    def _get_file_hash(self, file_path: str) -> str:
        """Generate hash of file for caching."""
        with open(file_path, 'rb') as f:
            return hashlib.md5(f.read()).hexdigest()
    
    def _is_text_meaningful(self, text: str, min_chars: int = 100) -> bool:
        """Check if extracted text is meaningful (not just garbage characters)."""
        if len(text.strip()) < min_chars:
            return False
        
        # Check if text has reasonable word-like patterns
        # Garbage OCR often has many special chars or very short "words"
        words = text.split()
        if len(words) < 10:
            return False
        
        # Check average word length (garbage tends to have very short or very long "words")
        avg_word_len = sum(len(w) for w in words) / len(words)
        if avg_word_len < 2 or avg_word_len > 20:
            return False
        
        # Check ratio of alphanumeric to total characters
        alnum_count = sum(1 for c in text if c.isalnum())
        if len(text) > 0 and alnum_count / len(text) < 0.5:
            return False
        
        return True
    
    def _extract_pdf_local(self, file_path: str) -> Tuple[str, bool]:
        """
        Extract text from PDF using pypdf (FREE).
        Returns (text, success_flag).
        """
        try:
            text_parts = []
            with open(file_path, 'rb') as file:
                reader = pypdf.PdfReader(file)
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text.strip())
            
            full_text = "\n\n".join(text_parts)
            
            # Consider successful if we got meaningful text
            if self._is_text_meaningful(full_text):
                return full_text, True
            return full_text, False
            
        except Exception as e:
            print(f"Local PDF extraction failed: {e}")
            return "", False
    
    def _extract_pdf_with_ocr(self, file_path: str) -> Tuple[str, bool]:
        """
        Extract text from scanned PDF using OCR (pytesseract).
        Requires: pip install pdf2image pytesseract
        And Tesseract OCR installed on system.
        Returns (text, success_flag).
        """
        if not PDF2IMAGE_AVAILABLE or not TESSERACT_AVAILABLE:
            return "", False
        
        try:
            # Convert PDF pages to images
            images = convert_from_path(file_path, dpi=200)  # type: ignore
            
            text_parts = []
            for i, image in enumerate(images):
                # Run OCR on each page image
                page_text = pytesseract.image_to_string(image)  # type: ignore
                if page_text and page_text.strip():
                    text_parts.append(f"--- Page {i+1} ---\n{page_text.strip()}")
            
            full_text = "\n\n".join(text_parts)
            
            if self._is_text_meaningful(full_text, min_chars=50):
                return full_text, True
            return full_text, False
            
        except Exception as e:
            print(f"OCR extraction failed: {e}")
            return "", False
    
    def _extract_pdf_with_ai_vision(self, file_path: str) -> str:
        """
        Extract text from scanned PDF using AI vision capabilities.
        Uses Gemini's vision model to read images.
        This costs money but works without Tesseract.
        """
        from app.services.unified_ai_service import ai_service
        
        if not ai_service:
            raise Exception("AI service not available")
        
        # Check if we have Gemini (which supports vision)
        if not ai_service.gemini_client:
            raise Exception("Gemini required for vision-based OCR")
        
        try:
            # Convert PDF to images
            if not PDF2IMAGE_AVAILABLE:
                raise Exception("pdf2image required for AI vision OCR. Install with: pip install pdf2image")
            
            images = convert_from_path(file_path, dpi=150)  # Lower DPI to reduce size  # type: ignore
            
            # Limit pages for cost
            max_pages = 10
            if len(images) > max_pages:
                print(f"Warning: PDF has {len(images)} pages, processing first {max_pages} only")
                images = images[:max_pages]
            
            text_parts = []
            
            for i, image in enumerate(images):
                # Convert PIL image to bytes
                import io
                import base64
                
                img_buffer = io.BytesIO()
                image.save(img_buffer, format='PNG')
                img_bytes = img_buffer.getvalue()
                img_b64 = base64.b64encode(img_bytes).decode('utf-8')
                
                # Use Gemini vision to extract text
                try:
                    import google.generativeai as genai  # type: ignore
                    
                    # Create image part for Gemini
                    image_part = {
                        "mime_type": "image/png",
                        "data": img_b64
                    }
                    
                    response = ai_service.gemini_client.generate_content([
                        "Extract ALL text from this image exactly as written. Include headings, paragraphs, lists, and any visible text. Return only the extracted text, no commentary.",
                        image_part
                    ])
                    
                    if response.candidates and response.candidates[0].content.parts:
                        page_text = response.candidates[0].content.parts[0].text
                        if page_text:
                            text_parts.append(f"--- Page {i+1} ---\n{page_text.strip()}")
                            ai_service.token_usage.add_usage(500, len(page_text) // 4)  # Rough estimate
                            
                except Exception as page_error:
                    print(f"Failed to process page {i+1}: {page_error}")
                    continue
            
            if not text_parts:
                raise Exception("No text could be extracted from any page")
            
            return "\n\n".join(text_parts)
            
        except Exception as e:
            raise Exception(f"AI vision OCR failed: {e}")
    
    def _extract_docx_local(self, file_path: str) -> Tuple[str, bool]:
        """
        Extract text from DOCX using python-docx (FREE).
        Returns (text, success_flag).
        """
        try:
            doc = Document(file_path)
            text_parts = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text.strip())
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_parts.append(row_text)
            
            full_text = "\n\n".join(text_parts)
            
            if len(full_text.strip()) > 50:
                return full_text, True
            return full_text, False
            
        except Exception as e:
            print(f"Local DOCX extraction failed: {e}")
            return "", False
    
    def _extract_txt(self, file_path: str) -> str:
        """Extract text from TXT file (always free and reliable)."""
        encodings = ['utf-8', 'latin-1', 'cp1252', 'ascii']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    return file.read().strip()
            except (UnicodeDecodeError, Exception):
                continue
        
        raise Exception("Failed to read TXT file with any encoding")
    
    def _extract_with_ai_fallback(self, file_path: str, file_type: str) -> str:
        """
        Fall back to AI extraction for complex documents (scanned PDFs, etc.)
        This costs money but handles edge cases.
        """
        from app.services.unified_ai_service import ai_service
        
        if not ai_service:
            raise Exception("AI service not available for fallback extraction")
        
        # Read file and convert to format AI can process
        with open(file_path, 'rb') as f:
            content = f.read()
        
        import base64
        file_b64 = base64.b64encode(content).decode('utf-8')
        
        # Only use AI for reasonable file sizes
        if len(file_b64) > 500000:  # ~375KB file
            raise Exception("File too large for AI extraction. Try a smaller file.")
        
        system = "Extract all text from this document. Return only the extracted text."
        user = f"Extract text from this {file_type.upper()} document (base64):\n{file_b64}"
        
        return ai_service._get_ai_response(system, user, temperature=0.2, max_tokens=4000)
    
    def extract_text(self, file_path: str, file_extension: str) -> str:
        """
        Extract text from file with smart fallback strategy:
        1. Check cache first
        2. Try local extraction (FREE)
        3. For PDFs: Try OCR with pytesseract (FREE if installed)
        4. For PDFs: Fall back to AI Vision (Gemini) as last resort
        """
        ext = file_extension.lower().lstrip('.')
        
        # Check cache
        file_hash = self._get_file_hash(file_path)
        if file_hash in self._text_cache:
            return self._text_cache[file_hash]
        
        extracted_text = ""
        
        if ext == 'txt':
            # TXT is always local
            extracted_text = self._extract_txt(file_path)
        
        elif ext == 'pdf':
            # Step 1: Try local text extraction (for regular PDFs)
            text, success = self._extract_pdf_local(file_path)
            if success:
                print("PDF text extraction successful (local)")
                extracted_text = text
            else:
                print("Local PDF extraction insufficient, trying OCR...")
                
                # Step 2: Try OCR with pytesseract (FREE)
                ocr_text, ocr_success = self._extract_pdf_with_ocr(file_path)
                if ocr_success:
                    print("OCR extraction successful (pytesseract)")
                    extracted_text = ocr_text
                else:
                    print("Pytesseract OCR failed or unavailable, trying AI Vision...")
                    
                    # Step 3: Try AI Vision (Gemini) as last resort
                    try:
                        extracted_text = self._extract_pdf_with_ai_vision(file_path)
                        print("AI Vision extraction successful (Gemini)")
                    except Exception as e:
                        print(f"AI Vision failed: {e}")
                        # Return whatever we got, even if incomplete
                        if ocr_text:
                            extracted_text = ocr_text
                        elif text:
                            extracted_text = text
                        else:
                            raise Exception(
                                f"Failed to extract text from scanned PDF. "
                                f"For best results, install Tesseract OCR and run: pip install pdf2image pytesseract. "
                                f"Error: {e}"
                            )
        
        elif ext in ['doc', 'docx']:
            # Try local first
            text, success = self._extract_docx_local(file_path)
            if success:
                extracted_text = text
            else:
                # DOCX shouldn't need OCR typically, but provide helpful error
                if text:
                    extracted_text = text
                else:
                    raise Exception("Failed to extract text from DOCX. The file may be corrupted.")
        
        else:
            raise ValueError(f"Unsupported file type: {ext}")
        
        # Cache the result
        if extracted_text:
            self._text_cache[file_hash] = extracted_text
        
        return extracted_text
        
        # Cache the result
        if extracted_text:
            self._text_cache[file_hash] = extracted_text
        
        return extracted_text
    
    def clear_cache(self):
        """Clear the text extraction cache."""
        self._text_cache.clear()
    
    def get_cache_stats(self) -> dict:
        """Get cache statistics."""
        return {
            "cached_files": len(self._text_cache),
            "cache_size_chars": sum(len(t) for t in self._text_cache.values())
        }
    
    def get_ocr_capabilities(self) -> dict:
        """Check what OCR capabilities are available."""
        return {
            "pdf2image_available": PDF2IMAGE_AVAILABLE,
            "pytesseract_available": TESSERACT_AVAILABLE,
            "local_ocr_ready": PDF2IMAGE_AVAILABLE and TESSERACT_AVAILABLE,
            "ai_vision_available": True,  # Always available if AI service configured
            "recommendation": (
                "Local OCR ready (FREE)" if (PDF2IMAGE_AVAILABLE and TESSERACT_AVAILABLE)
                else "Install for free OCR: pip install pdf2image pytesseract (and Tesseract OCR)"
            )
        }


# Global instance
file_processing_service = OptimizedFileService()
